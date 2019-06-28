import docx
import pandas as pd
from docx.api import Document
document='E:/Text_model/goibibo.docx'
def readtxt(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

print (readtxt(document))

document = Document(document)
table = document.tables[0]

data = []

keys = None
for i, row in enumerate(table.rows):
    text = (cell.text for cell in row.cells)

    if i == 0:
        keys = tuple(text)
        continue
    row_data = dict(zip(keys, text))
    data.append(row_data)
    print (data)

df = pd.DataFrame(data)
#df=df.T
df.to_csv('E:/Text_model/hhahah.csv')

